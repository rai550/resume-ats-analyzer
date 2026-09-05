import streamlit as st
import json
import os
import re
from io import BytesIO

from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document

st.set_page_config(page_title="AI Resume ATS Analyzer", page_icon="📄", layout="wide")

MODEL_NAME = "gemini-3.6-flash"


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            text = " | ".join(cell.text.strip() for cell in row.cells)
            if text.strip():
                parts.append(text)
    return "\n".join(parts)


def extract_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    filename = uploaded_file.name.lower()
    if filename.endswith(".pdf"):
        return extract_pdf(data)
    if filename.endswith(".docx"):
        return extract_docx(data)
    if filename.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


def clean_json(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


def get_api_key():
    try:
        key = st.secrets.get("GEMINI_API_KEY")
    except Exception:
        key = None
    return key or os.getenv("GEMINI_API_KEY")


def analyze_resume(resume_text: str, job_description: str) -> dict:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError(
            "GEMINI_API_KEY is not configured. Add it to Streamlit Secrets "
            "or your environment variables."
        )

    client = genai.Client(api_key=api_key)
    jd = job_description.strip() or (
        "No specific job description was provided. "
        "Evaluate the resume for general ATS readiness."
    )

    prompt = f"""
You are an expert ATS resume analyst, technical recruiter, career coach,
and resume optimization specialist.

Analyze the resume against the provided job description.

This is an explainable ESTIMATED ATS READINESS SCORE, NOT the exact score
of any proprietary ATS such as Workday, Greenhouse, Lever, or Taleo.

Evaluate keyword alignment, job relevance, standard sections, skills,
experience, quantified achievements, contact information, ATS-friendly
structure, parsing risks, and overall quality.

Return ONLY valid JSON using exactly this structure:
{{
  "ats_score": 0,
  "score_label": "Excellent|Good|Needs Improvement|Poor",
  "summary": "Short overall assessment.",
  "keyword_match": {{"matched": [], "missing": [], "partial": []}},
  "section_scores": {{
    "contact_information": 0,
    "professional_summary": 0,
    "experience": 0,
    "skills": 0,
    "education": 0,
    "formatting_ats_readiness": 0
  }},
  "strengths": [],
  "improvements": [
    {{"priority": "High|Medium|Low", "issue": "",
      "recommendation": "", "example": ""}}
  ],
  "ats_checklist": {{
    "standard_headings": true,
    "keyword_alignment": true,
    "quantified_achievements": true,
    "simple_ats_friendly_format": true,
    "contact_information_present": true,
    "no_obvious_parsing_issues": true
  }}
}}

Scoring: 90-100 Excellent; 75-89 Good; 60-74 Needs Improvement; 0-59 Poor.

Do NOT invent work experience, skills, education, certifications, employers,
job titles, achievements, or metrics. Recommendations must be based on the
resume or job description. Any new wording must clearly be an example.

JOB DESCRIPTION:
{jd}

RESUME:
{resume_text}
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(response_mime_type="application/json"),
    )

    if not response.text:
        raise ValueError("Gemini returned an empty response.")

    result = clean_json(response.text)
    if not isinstance(result, dict):
        raise ValueError("Gemini returned an unexpected response.")

    try:
        score = int(result.get("ats_score", 0))
    except (TypeError, ValueError):
        score = 0
    result["ats_score"] = max(0, min(100, score))
    return result


def score_icon(score: int) -> str:
    if score >= 90:
        return "🟢"
    if score >= 75:
        return "🟡"
    if score >= 60:
        return "🟠"
    return "🔴"


st.title("📄 AI Resume ATS Analyzer")
st.caption(
    "Upload your resume and get an explainable ATS-readiness score, "
    "keyword analysis, strengths, and improvement recommendations."
)

with st.sidebar:
    st.header("⚙️ Settings")
    st.info("The Gemini API key is read securely from Streamlit Secrets as GEMINI_API_KEY.")
    st.markdown("### Supported Files")
    st.markdown("- PDF\n- DOCX\n- TXT")
    st.markdown("### AI Model")
    st.code(MODEL_NAME)

uploaded_file = st.file_uploader(
    "📤 Upload your resume",
    type=["pdf", "docx", "txt"],
    help="Upload a text-based PDF, DOCX, or TXT resume.",
)

job_description = st.text_area(
    "💼 Paste the Job Description",
    height=240,
    placeholder="Paste the target job description here for a more useful analysis.",
)

if uploaded_file:
    try:
        resume_text = extract_text(uploaded_file)

        if not resume_text.strip():
            st.error("I couldn't extract readable text from this file. Try a text-based PDF or DOCX.")
            st.stop()

        word_count = len(resume_text.split())
        st.success(f"Resume loaded successfully: {uploaded_file.name} • {word_count:,} words")

        with st.expander("👀 Preview Extracted Resume Text"):
            st.text(resume_text[:12000])

        if st.button("🔍 Analyze Resume", type="primary", use_container_width=True):
            with st.spinner("Gemini Flash is analyzing your resume..."):
                try:
                    st.session_state["analysis"] = analyze_resume(resume_text, job_description)
                except Exception as error:
                    st.error(f"Analysis failed: {error}")
    except Exception as error:
        st.error(f"Could not read the resume: {error}")

analysis = st.session_state.get("analysis")

if analysis:
    st.divider()

    score = analysis.get("ats_score", 0)
    st.metric("Estimated ATS Score", f"{score}/100")

    c1, c2 = st.columns([1, 3])
    with c1:
        st.subheader(f"{score_icon(score)} {analysis.get('score_label', 'Assessment')}")
    with c2:
        st.write(analysis.get("summary", ""))
        st.info("This is an AI-based ATS-readiness estimate, not the exact score from a commercial ATS.")

    st.subheader("🔑 Keyword Match")
    keywords = analysis.get("keyword_match", {})
    k1, k2, k3 = st.columns(3)
    with k1:
        st.markdown("### 🟢 Matched")
        st.write(", ".join(keywords.get("matched", [])) or "None identified.")
    with k2:
        st.markdown("### 🔴 Missing")
        st.write(", ".join(keywords.get("missing", [])) or "None identified.")
    with k3:
        st.markdown("### 🟡 Partial")
        st.write(", ".join(keywords.get("partial", [])) or "None identified.")

    st.subheader("📊 Resume Section Scores")
    for section, value in analysis.get("section_scores", {}).items():
        try:
            value = max(0, min(100, int(value)))
        except (TypeError, ValueError):
            value = 0
        st.write(f"**{section.replace('_', ' ').title()}: {value}/100**")
        st.progress(value / 100)

    st.subheader("💪 Resume Strengths")
    strengths = analysis.get("strengths", [])
    if strengths:
        for item in strengths:
            st.markdown(f"- {item}")
    else:
        st.write("No specific strengths were returned.")

    st.subheader("🚀 Recommended Improvements")
    improvements = analysis.get("improvements", [])
    if improvements:
        for item in improvements:
            st.markdown(f"### {item.get('priority', 'Medium')} — {item.get('issue', '')}")
            st.write(item.get("recommendation", ""))
            if item.get("example"):
                st.caption(f"Example: {item['example']}")
    else:
        st.write("No improvement recommendations were returned.")

    st.subheader("✅ ATS Checklist")
    for key, value in analysis.get("ats_checklist", {}).items():
        st.checkbox(key.replace("_", " ").title(), value=bool(value), disabled=True)

    st.divider()
    st.download_button(
        "⬇️ Download Analysis JSON",
        data=json.dumps(analysis, indent=2, ensure_ascii=False),
        file_name="resume_ats_analysis.json",
        mime="application/json",
        use_container_width=True,
    )
else:
    st.info("👆 Upload your resume, optionally paste a job description, and click Analyze Resume.")
